"""
文件文本提取模块 (File Text Extraction)

【安全检测全局说明 - 文件解析风险】
1. 文件解析是引发拒绝服务 (DoS) 和远程代码执行 (RCE) 的高危区域。
2. 必须在 API 路由层（如 FastAPI 的 UploadFile）严格限制上传文件的大小（如 MAX_UPLOAD_SIZE = 10MB），
   防止 Zip Bomb (压缩包炸弹) 或超大文件导致内存耗尽 (OOM)。
3. 【高危警告】.doc 格式依赖本地 Microsoft Word COM 组件解析。Word 解析器历史上存在大量 RCE 漏洞，
   强烈建议在生产环境（尤其是 Linux/Docker）中禁用 .doc 支持，或改用隔离的沙箱环境（如 LibreOffice headless）。
"""

from __future__ import annotations

import io
import logging
import os
import tempfile
from pathlib import Path
from typing import Final

from fastapi import HTTPException, status

# 【规范说明】使用模块级 logger，避免使用 basicConfig 污染全局日志配置
logger = logging.getLogger(__name__)

# ==========================================
# 1. 常量配置 (Constants)
# ==========================================

SUPPORTED_EXTENSIONS: Final[set[str]] = {".txt", ".docx", ".doc"}

# 【安全说明】按优先级排列的文本编码。utf-8-sig 用于处理带 BOM 头的 Windows 记事本文件。
TXT_ENCODINGS: Final[tuple[str, ...]] = ("utf-8-sig", "utf-8", "gb18030", "gbk")


# ==========================================
# 2. TXT 文本解析 (TXT Parsing)
# ==========================================

def _decode_text_bytes(raw: bytes) -> str:
    """
    尝试多种编码将字节流解码为字符串。

    【安全检测说明】
    Python 内置的 decode 方法对恶意字节流有较好的容错性，不会引发 ReDoS。
    但需注意，如果 raw 极大（如几百 MB），多次 decode 尝试会导致 CPU 飙升。
    因此，必须在上传接口层限制文件大小。
    """
    for encoding in TXT_ENCODINGS:
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="TXT 文件编码无法识别，请保存为 UTF-8 或 GBK 格式"
    )


# ==========================================
# 3. DOCX 文本解析 (DOCX Parsing)
# ==========================================

def _extract_docx_text(raw: bytes) -> str:
    """
    从 .docx 文件中提取文本（包括段落和表格）。

    【安全检测说明 - XXE 与 实体膨胀防御】
    .docx 本质上是包含 XML 文件的 ZIP 压缩包。
    1. 必须确保底层 XML 解析器防御了 XXE (XML External Entity) 攻击。
       python-docx 默认使用 lxml，建议全局注入 defusedxml 以彻底免疫 XXE 和 Billion Laughs 攻击。
    2. 恶意构造的 docx 可能包含数百万个 XML 节点，导致解析时 CPU/内存耗尽。
    """
    try:
        # 【安全加固】强制使用 defusedxml 替换标准 xml 解析器，防御 XXE 攻击
        try:
            import defusedxml
            defusedxml.defuse()
        except ImportError:
            logger.warning("defusedxml 未安装，docx 解析可能存在 XXE 风险。请执行: pip install defusedxml")

        from docx import Document
    except ImportError as exc:
        logger.error("缺少 docx 解析库: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="服务器缺少文档解析依赖"
        ) from exc

    try:
        document = Document(io.BytesIO(raw))
    except Exception as exc:
        logger.warning("docx 文件结构损坏或包含恶意 XML: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="DOCX 文件损坏或格式不合法"
        ) from exc

    # 提取段落文本
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    if paragraphs:
        return "\n".join(paragraphs)

    # 降级提取表格文本（若段落为空）
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append("\t".join(cells))

    return "\n".join(table_lines)


# ==========================================
# 4. DOC 文本解析 (Legacy DOC Parsing via COM)
# ==========================================

def _extract_doc_text(raw: bytes) -> str:
    """
    使用 Windows COM 组件调用本地 Microsoft Word 解析旧版 .doc 文件。

    【安全检测核心警告 - RCE 与 DoS 风险】
    1. RCE 风险：将不受信任的二进制数据写入磁盘并交给本地 Word 解析，极易触发 Word 解析器漏洞导致远程代码执行。
    2. DoS 风险：COM 组件调用容易因文件损坏导致 Word 进程挂起（僵尸进程），最终耗尽服务器资源。
    3. 跨平台限制：此功能仅限 Windows 环境，且必须安装 Microsoft Word。
    """
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="当前环境不支持 .doc 文件解析（需 Windows 及 Word 组件），请转换为 .docx 格式后重试",
        ) from exc

    temp_path: str | None = None
    word_app = None

    try:
        # 1. 安全写入临时文件
        # 【安全说明】使用 tempfile 确保文件名不可预测，防止路径遍历和临时文件竞争条件 (TOCTOU)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(raw)
            temp_path = temp_file.name

        # 2. 初始化 COM 环境
        pythoncom.CoInitialize()

        # 3. 启动 Word 进程
        word_app = DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0  # 禁用所有弹窗警告，防止进程阻塞

        # 4. 打开并提取文本
        # ReadOnly=1 防止恶意宏修改文件，ConfirmConversions=False 防止弹窗
        doc = word_app.Documents.Open(temp_path, ReadOnly=1, ConfirmConversions=False)
        try:
            text = doc.Content.Text
        finally:
            doc.Close(False)

        return text.strip()

    except HTTPException:
        raise
    except Exception as exc:
        # 【安全说明 - 信息泄露防护】
        # 严禁将 exc 的具体信息（可能包含内部路径、COM 错误码）直接返回给客户端
        logger.error("COM 解析 .doc 文件失败: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=".doc 文件解析失败，文件可能已损坏或包含不支持的格式"
        ) from exc
    finally:
        # 5. 严格的资源清理 (防止僵尸进程和临时文件泄露)
        if word_app is not None:
            try:
                word_app.Quit()
            except Exception as e:
                logger.warning("Word COM 进程强制退出失败: %s", e)

        try:
            pythoncom.CoUninitialize()
        except Exception as e:
            logger.debug("COM 组件反初始化异常 (可忽略): %s", e)

        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except OSError as e:
                logger.warning("临时文件清理失败 [%s]: %s", temp_path, e)


# ==========================================
# 5. 统一路由分发 (Routing & Dispatch)
# ==========================================

def extract_text_from_file(filename: str, raw: bytes) -> str:
    """
    根据文件扩展名路由到对应的文本提取器。

    Args:
        filename: 原始文件名（用于判断后缀）。
        raw: 文件的原始字节流。

    Returns:
        提取出的纯文本内容。

    Raises:
        HTTPException: 文件格式不支持或解析失败时抛出。
    """
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件格式。仅允许: {', '.join(SUPPORTED_EXTENSIONS)}",
        )

    if suffix == ".txt":
        return _decode_text_bytes(raw).strip()
    if suffix == ".docx":
        return _extract_docx_text(raw).strip()
    if suffix == ".doc":
        return _extract_doc_text(raw).strip()

    # 理论上不会执行到这里，但为了类型安全和防御性编程保留
    raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="内部路由错误")